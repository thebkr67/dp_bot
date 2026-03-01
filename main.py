import os
import io
import re
import json
import base64
import logging
import asyncio
import httpx
import time
from dataclasses import dataclass
from typing import Optional, Tuple

from dotenv import load_dotenv

from aiogram import Bot, Dispatcher, F
from aiogram.types import Message, FSInputFile, CallbackQuery, ReactionTypeEmoji
from aiogram.enums import ChatAction
from aiogram.filters import Command
from aiogram.utils.keyboard import InlineKeyboardBuilder

from openai import OpenAI

from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook, Workbook
from PIL import Image


# ---------- setup ----------
load_dotenv()

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
SERPER_API_KEY = os.getenv("SERPER_API_KEY")

DEEPAI_API_KEY = os.getenv("DEEPAI_API_KEY")
DEEPAI_TIMEOUT_SEC = int(os.getenv("DEEPAI_TIMEOUT_SEC", "120"))
IMG_FALLBACK = os.getenv("IMG_FALLBACK", "openai").lower()  # openai/none
ENHANCE_DEFAULT_SCALE = int(os.getenv("ENHANCE_DEFAULT_SCALE", "2"))  # 2 or 4

# --- Pika Labs (image -> video) via fal.ai ---
FAL_KEY = os.getenv("FAL_KEY") or os.getenv("FAL_API_KEY")
FAL_QUEUE_BASE = os.getenv("FAL_QUEUE_BASE", "https://queue.fal.run")
PIKA_IMAGE2VIDEO_MODEL = os.getenv("PIKA_IMAGE2VIDEO_MODEL", "fal-ai/pika/v2.2/image-to-video")
PIKA_DEFAULT_RESOLUTION = os.getenv("PIKA_DEFAULT_RESOLUTION", "720p")  # 720p or 1080p
PIKA_DEFAULT_DURATION = int(os.getenv("PIKA_DEFAULT_DURATION", "5"))  # 5 or 10
PIKA_POLL_INTERVAL_SEC = int(os.getenv("PIKA_POLL_INTERVAL_SEC", "5"))
PIKA_TASK_TIMEOUT_SEC = int(os.getenv("PIKA_TASK_TIMEOUT_SEC", "600"))

if not DEEPAI_API_KEY:
    logging.warning("DEEPAI_API_KEY is not set (DeepAI features disabled)")


if not TELEGRAM_BOT_TOKEN:
    raise RuntimeError("TELEGRAM_BOT_TOKEN is not set")
if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY is not set")

logging.basicConfig(level=logging.INFO)

client = OpenAI(api_key=OPENAI_API_KEY)
dp = Dispatcher()

TEXT_MODEL = "gpt-4o-mini"
VISION_MODEL = "gpt-4o-mini"

MAX_FILE_BYTES = 12 * 1024 * 1024  # 12 MB

# ---------- autosearch settings ----------
SEARCH_TRIGGERS = [
    "найди", "поищи", "поиск", "в интернете", "гугл", "google",
    "ссылк", "источник", "пруф", "докажи", "где написано",
    "актуаль", "сейчас", "на сегодня", "последн", "свеж",
    "новост", "цена", "стоимость", "тариф", "курс", "ставк",
    "правила", "регламент", "инструкция", "политика", "обновил",
    "railway", "serper", "aiogram", "wildberries", "ozon"
]

SEARCH_MIN_LEN = 18               # минимальная длина сообщения для автопоиска
SEARCH_RESULTS_NUM = 5            # сколько результатов подмешивать
SEARCH_COOLDOWN_SEC = 12          # 1 поиск / 12 сек на пользователя
SEARCH_CACHE_TTL_SEC = 300        # кэш выдачи на 5 минут


class Reference:
    def __init__(self) -> None:
        self.response = ""


reference = Reference()


def clear_past():
    reference.response = ""


# ---------- last file storage ----------
@dataclass
class LastFile:
    filename: str
    ext: str
    mime: str
    data: bytes


last_files: dict[int, LastFile] = {}  # key = telegram user_id

# ---------- last image storage ----------
ENHANCE_CB = "enhance:last"
RETROUCH_CB = "retouch:last"
VIDEO_CB = "video:last"
last_images: dict[int, bytes] = {}  # key = telegram user_id
pending_video_prompt: dict[int, bool] = {}  # user_id -> ждём текст-подсказку для видео


# ---------- autosearch state ----------
_last_search_at: dict[int, float] = {}
_search_cache: dict[str, tuple[float, list[dict]]] = {}


def should_autosearch(text: str) -> bool:
    """
    Решаем, нужен ли автопоиск.
    """
    t = (text or "").strip().lower()
    if not t:
        return False
    if t.startswith("/"):
        return False  # команды не трогаем
    if len(t) < SEARCH_MIN_LEN:
        return False
    return any(k in t for k in SEARCH_TRIGGERS)


def can_search_now(user_id: int) -> bool:
    now = time.time()
    last = _last_search_at.get(user_id, 0.0)
    if now - last < SEARCH_COOLDOWN_SEC:
        return False
    _last_search_at[user_id] = now
    return True


def _cache_key(query: str, num: int) -> str:
    q = re.sub(r"\s+", " ", (query or "").strip().lower())
    return f"{q}::num={num}"


async def serper_search(query: str, num: int = 5) -> list[dict]:
    """
    Возвращает список результатов: title, link, snippet
    """
    if not SERPER_API_KEY:
        raise RuntimeError("SERPER_API_KEY is not set")

    url = "https://google.serper.dev/search"
    headers = {
        "X-API-KEY": SERPER_API_KEY,
        "Content-Type": "application/json",
    }
    payload = {"q": query, "num": max(1, min(num, 10))}

    async with httpx.AsyncClient(timeout=20) as client_http:
        r = await client_http.post(url, headers=headers, json=payload)
        r.raise_for_status()
        data = r.json()

    results = []
    for item in (data.get("organic") or [])[:num]:
        results.append({
            "title": item.get("title", ""),
            "link": item.get("link", ""),
            "snippet": item.get("snippet", ""),
        })
    return results


async def serper_search_cached(query: str, num: int = 5) -> list[dict]:
    """
    Кэшируем поиск, чтобы не жечь лимиты.
    """
    key = _cache_key(query, num)
    now = time.time()

    if key in _search_cache:
        ts, cached = _search_cache[key]
        if now - ts < SEARCH_CACHE_TTL_SEC:
            return cached

    results = await serper_search(query, num=num)
    _search_cache[key] = (now, results)
    return results


def format_search_results(results: list[dict]) -> str:
    text = "🔎 Результаты поиска:\n\n"
    for i, r in enumerate(results, start=1):
        text += f"{i}) {r.get('title','')}\n{r.get('link','')}\n{r.get('snippet','')}\n\n"
    return text.strip()


def format_results_for_prompt(results: list[dict]) -> str:
    # компактный блок для промпта модели
    lines = []
    for r in results:
        title = (r.get("title") or "").strip()
        snippet = (r.get("snippet") or "").strip()
        link = (r.get("link") or "").strip()
        if title or snippet or link:
            lines.append(f"- {title}: {snippet} ({link})")
    return "\n".join(lines).strip()


# ---------- helpers ----------
def _ext(filename: Optional[str]) -> str:
    if not filename:
        return ""
    m = re.search(r"(\.[a-zA-Z0-9]+)$", filename)
    return (m.group(1) if m else "").lower()


def _safe_truncate(s: str, max_chars: int = 45_000) -> str:
    s = s or ""
    if len(s) <= max_chars:
        return s
    return s[:max_chars] + "\n\n...[обрезано]..."


def _extract_text_from_pdf(data: bytes, max_pages: int = 10) -> str:
    reader = PdfReader(io.BytesIO(data))
    chunks = []
    for page in reader.pages[:max_pages]:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            chunks.append("")
    return "\n".join(chunks).strip()


def _extract_text_from_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    paras = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paras).strip()


def _extract_tsv_preview_from_xlsx(data: bytes, max_rows: int = 100, max_cols: int = 25) -> str:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    ws = wb.active
    rows_out = []
    for r_i, row in enumerate(ws.iter_rows(values_only=True), start=1):
        if r_i > max_rows:
            break
        row = row[:max_cols]
        row_s = ["" if v is None else str(v) for v in row]
        rows_out.append("\t".join(row_s).rstrip())
    wb.close()
    return "\n".join(rows_out).strip()


def _extract_text_from_plain(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("cp1251", errors="replace")


def _detect_image_mime_from_bytes(data: bytes) -> str:
    img = Image.open(io.BytesIO(data))
    fmt = (img.format or "PNG").upper()
    if fmt == "JPEG":
        return "image/jpeg"
    if fmt == "WEBP":
        return "image/webp"
    return "image/png"


def _to_data_uri(image_bytes: bytes) -> str:
    """Encode image bytes to a data URI for Runway API input."""
    mime = _detect_image_mime_from_bytes(image_bytes)
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    return f"data:{mime};base64,{b64}"



async def _download_telegram_file(bot: Bot, file_id: str) -> Tuple[bytes, str]:
    file = await bot.get_file(file_id)
    buf = io.BytesIO()
    await bot.download_file(file.file_path, destination=buf)
    return buf.getvalue(), (file.file_path or "")


async def _ask_openai_text(system: str, user: str) -> str:
    resp = client.chat.completions.create(
        model=TEXT_MODEL,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    )
    return resp.choices[0].message.content or ""


async def _ask_openai_vision(prompt: str, image_bytes: bytes) -> str:
    mime = _detect_image_mime_from_bytes(image_bytes)
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    data_url = f"data:{mime};base64,{b64}"

    r = client.responses.create(
        model=VISION_MODEL,
        input=[
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": prompt},
                    {"type": "input_image", "image_url": data_url},
                ],
            }
        ],
    )
    return (r.output_text or "").strip()


def _save_bytes_to_tmp(filename: str, data: bytes) -> str:
    safe = re.sub(r"[^a-zA-Z0-9._-]+", "_", filename)[:120]
    path = f"/tmp/{safe}"
    with open(path, "wb") as f:
        f.write(data)
    return path


# ---------- file building ----------
def _build_docx_from_text(text: str) -> bytes:
    doc = DocxDocument()
    for para in (text or "").splitlines():
        doc.add_paragraph(para)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _build_xlsx_from_tsv(tsv: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    for r_i, line in enumerate((tsv or "").splitlines(), start=1):
        cols = line.split("\t")
        for c_i, v in enumerate(cols, start=1):
            ws.cell(row=r_i, column=c_i, value=v)
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ---------- edit engines ----------
async def _edit_text_like(original_text: str, instructions: str) -> str:
    system = (
        "Ты редактор. Тебе дают исходный текст и инструкцию, что изменить. "
        "Верни ТОЛЬКО финальную исправленную версию текста без комментариев."
    )
    user = (
        f"ИНСТРУКЦИЯ:\n{instructions}\n\n"
        f"ИСТОЧНИК:\n-----\n{_safe_truncate(original_text)}\n-----\n"
    )
    return await _ask_openai_text(system, user)


async def _edit_docx_bytes(data: bytes, instructions: str) -> Tuple[bytes, str]:
    src = _extract_text_from_docx(data)
    edited = await _edit_text_like(src, instructions)
    return _build_docx_from_text(edited), "docx"


async def _edit_xlsx_bytes(data: bytes, instructions: str) -> Tuple[bytes, str]:
    preview = _extract_tsv_preview_from_xlsx(data, max_rows=120, max_cols=25)
    system = (
        "Ты ассистент по таблицам. Тебе дают TSV-таблицу (таб-разделители) и инструкцию, что изменить. "
        "Верни ТОЛЬКО итоговый TSV (табами, строки переносами), без пояснений, без markdown."
    )
    user = (
        f"ИНСТРУКЦИЯ:\n{instructions}\n\n"
        f"ТАБЛИЦА TSV:\n-----\n{preview}\n-----\n"
        "Требования: сохраняй структуру таблицы, не добавляй лишних комментариев."
    )
    tsv = await _ask_openai_text(system, user)
    tsv = tsv.strip().strip("```").strip()
    return _build_xlsx_from_tsv(tsv), "xlsx"



# ---------- DeepAI image generation / enhance ----------


def _image_action_keyboard():
    kb = InlineKeyboardBuilder()
    kb.button(text="✨ Улучшить", callback_data=ENHANCE_CB)
    kb.button(text="🛍️ WB ретушь", callback_data=RETROUCH_CB)
    kb.button(text="🎥 Сделать видео", callback_data=VIDEO_CB)
    kb.adjust(2, 1)
    return kb.as_markup()


async def _openai_image_generate(prompt: str) -> bytes:
    res = client.images.generate(
        model="gpt-image-1",
        prompt=prompt,
    )
    b64 = res.data[0].b64_json
    return base64.b64decode(b64)


async def _openai_image_edit(image_bytes: bytes, prompt: str) -> bytes:
    mime = _detect_image_mime_from_bytes(image_bytes)
    img_file = io.BytesIO(image_bytes)
    img_file.name = "input.png" if mime == "image/png" else "input.jpg"
    res = client.images.edit(
        model="gpt-image-1",
        image=img_file,
        prompt=prompt,
    )
    b64 = res.data[0].b64_json
    return base64.b64decode(b64)


async def generate_image_from_text_with_fallback(prompt: str) -> bytes:
    try:
        return await generate_image_from_text(prompt)
    except Exception:
        if IMG_FALLBACK == "openai":
            logging.exception("DeepAI text2img failed, using OpenAI fallback")
            return await _openai_image_generate(prompt)
        raise


async def enhance_image_with_fallback(image_bytes: bytes, scale: int = 2) -> bytes:
    """scale: 2 or 4. DeepAI torch-srgan is ~2x; for 4x run twice."""
    scale = 4 if int(scale) == 4 else 2
    try:
        out = await enhance_image(image_bytes)
        if scale == 4:
            out = await enhance_image(out)
        return out
    except Exception:
        if IMG_FALLBACK == "openai":
            logging.exception("DeepAI enhance failed, using OpenAI fallback")
            prompt = "Улучши качество фото: резкость, детализация, шумоподавление, натуральные цвета. Без изменения смысла."
            out = await _openai_image_edit(image_bytes, prompt)
            if scale == 4:
                out = await _openai_image_edit(out, prompt)
            return out
        raise


async def retouch_for_wb(image_bytes: bytes) -> bytes:
    """AI ретушь для карточек WB: чище, контрастнее, без изменения товара."""
    prompt = (
        "Сделай ретушь фото товара для маркетплейса Wildberries: "
        "увеличь четкость и детализацию, убери шум и грязные оттенки, "
        "сделай фон максимально чистым и аккуратным (без лишних объектов), "
        "сохрани натуральные цвета и фактуру товара, "
        "не меняй форму и дизайн товара. "
        "Результат должен выглядеть как качественная студийная предметная съемка."
    )
    if IMG_FALLBACK == "openai":
        return await _openai_image_edit(image_bytes, prompt)
    # Best-effort: у DeepAI нет промпта в SRGAN, поэтому просто улучшение.
    return await enhance_image(image_bytes)


DEEPAI_TEXT2IMG_URL = "https://api.deepai.org/api/text2img"
DEEPAI_UPSCALE_URL = "https://api.deepai.org/api/torch-srgan"


async def generate_image_from_text(prompt: str) -> bytes:
    prompt = (prompt or "").strip()
    if not prompt:
        raise ValueError("Empty prompt")

    files = {"text": (None, prompt)}
    headers = {"api-key": DEEPAI_API_KEY}

    async with httpx.AsyncClient(timeout=DEEPAI_TIMEOUT_SEC) as client_http:
        r = await client_http.post(DEEPAI_TEXT2IMG_URL, headers=headers, files=files)
        r.raise_for_status()
        data = r.json()

    output_url = data.get("output_url")
    if not output_url:
        raise RuntimeError(f"DeepAI error: {data}")

    async with httpx.AsyncClient(timeout=DEEPAI_TIMEOUT_SEC) as client_http:
        img = await client_http.get(output_url)
        img.raise_for_status()
        return img.content


async def enhance_image(image_bytes: bytes) -> bytes:
    if not image_bytes:
        raise ValueError("No image")

    headers = {"api-key": DEEPAI_API_KEY}
    files = {"image": ("input.jpg", image_bytes, "image/jpeg")}

    async with httpx.AsyncClient(timeout=DEEPAI_TIMEOUT_SEC) as client_http:
        r = await client_http.post(DEEPAI_UPSCALE_URL, headers=headers, files=files)
        r.raise_for_status()
        data = r.json()

    output_url = data.get("output_url")
    if not output_url:
        raise RuntimeError(f"DeepAI error: {data}")

    async with httpx.AsyncClient(timeout=DEEPAI_TIMEOUT_SEC) as client_http:
        img = await client_http.get(output_url)
        img.raise_for_status()
        return img.content



# ---------------- Pika (fal.ai) image -> video ----------------
# Docs:
# - Model: fal-ai/pika/v2.2/image-to-video
# - Submit: POST https://queue.fal.run/<model_id> with JSON body matching schema
# - Status: GET  https://queue.fal.run/<model_id>/requests/<request_id>/status
# - Result: GET  https://queue.fal.run/<model_id>/requests/<request_id>
# Auth: Authorization: Key <FAL_KEY>
# Files: you can pass image_url as hosted URL or as base64 data URI.
# Sources:
# - Pika v2.2 schema + inputs/outputs: https://fal.ai/models/fal-ai/pika/v2.2/image-to-video/api
# - Queue API: https://docs.fal.ai/model-apis/model-endpoints/queue

def _fal_headers() -> dict:
    if not FAL_KEY:
        raise RuntimeError("FAL_KEY is not set")
    return {
        "Authorization": f"Key {FAL_KEY}",
        "Content-Type": "application/json",
    }


async def fal_queue_submit(model_id: str, payload: dict) -> dict:
    url = f"{FAL_QUEUE_BASE}/{model_id}"
    headers = _fal_headers()
    async with httpx.AsyncClient(timeout=60) as client_http:
        r = await client_http.post(url, headers=headers, json=payload)
        r.raise_for_status()
        return r.json()


async def fal_queue_status(model_id: str, request_id: str, *, logs: bool = False) -> dict:
    q = "?logs=1" if logs else ""
    url = f"{FAL_QUEUE_BASE}/{model_id}/requests/{request_id}/status{q}"
    headers = _fal_headers()
    async with httpx.AsyncClient(timeout=60) as client_http:
        r = await client_http.get(url, headers=headers)
        r.raise_for_status()
        return r.json()


async def fal_queue_result(model_id: str, request_id: str) -> dict:
    url = f"{FAL_QUEUE_BASE}/{model_id}/requests/{request_id}"
    headers = _fal_headers()
    async with httpx.AsyncClient(timeout=120) as client_http:
        r = await client_http.get(url, headers=headers)
        r.raise_for_status()
        return r.json()


async def fal_wait_for_result(model_id: str, request_id: str, *, timeout_sec: int = PIKA_TASK_TIMEOUT_SEC) -> dict:
    start = time.time()
    last_status = None
    while True:
        status_obj = await fal_queue_status(model_id, request_id, logs=False)
        status = status_obj.get("status")
        if status and status != last_status:
            last_status = status

        if status == "COMPLETED":
            return await fal_queue_result(model_id, request_id)

        # IN_QUEUE / IN_PROGRESS return 202 from status endpoint; json still has status field.
        if time.time() - start > timeout_sec:
            raise TimeoutError(f"fal.ai task timeout after {timeout_sec}s (last status: {status})")

        await asyncio.sleep(max(1, PIKA_POLL_INTERVAL_SEC))


async def fal_wait_for_result_by_urls(status_url: str, response_url: str, *, timeout_sec: int = PIKA_TASK_TIMEOUT_SEC) -> dict:
    """Wait for completion using the status_url/response_url returned by fal submit.
    This avoids the 'subpath' pitfall where status/result endpoints omit the subpath."""
    start = time.time()
    last_status = None
    headers = _fal_headers()

    async with httpx.AsyncClient(timeout=60) as client_http:
        while True:
            r = await client_http.get(status_url, headers=headers)
            # status endpoint returns 202 for IN_QUEUE/IN_PROGRESS, 200 for COMPLETED
            if r.status_code not in (200, 202):
                r.raise_for_status()
            status_obj = r.json()
            status = status_obj.get("status")
            if status and status != last_status:
                last_status = status

            if status == "COMPLETED":
                rr = await client_http.get(response_url, headers=headers)
                rr.raise_for_status()
                return rr.json()

            if time.time() - start > timeout_sec:
                raise TimeoutError(f"fal.ai task timeout after {timeout_sec}s (last status: {status})")

            await asyncio.sleep(max(1, PIKA_POLL_INTERVAL_SEC))



async def pika_image_bytes_to_video(
    image_bytes: bytes,
    *,
    prompt: str,
    resolution: str = PIKA_DEFAULT_RESOLUTION,
    duration: int = PIKA_DEFAULT_DURATION,
    negative_prompt: str = "",
    seed: int | None = None,
) -> tuple[str, bytes]:
    """Return (request_id, mp4_bytes) using fal-ai/pika/v2.2/image-to-video."""
    if not image_bytes:
        raise ValueError("No image bytes")
    payload = {
        "image_url": _to_data_uri(image_bytes),
        "prompt": prompt,
        "resolution": resolution,
        "duration": int(duration),
    }
    if negative_prompt:
        payload["negative_prompt"] = negative_prompt
    if seed is not None:
        payload["seed"] = int(seed)

    submit = await fal_queue_submit(PIKA_IMAGE2VIDEO_MODEL, payload)
    request_id = submit.get("request_id")
    status_url = submit.get("status_url")
    response_url = submit.get("response_url")
    if not request_id or not status_url or not response_url:
        raise RuntimeError(f"fal.ai submit failed: {submit}")

    result = await fal_wait_for_result_by_urls(status_url, response_url)
    # Result shape typically: {"video": {"url": "...mp4"}}
    video_url = ((result or {}).get("video") or {}).get("url")
    if not video_url:
        raise RuntimeError(f"fal.ai result missing video url: {result}")

    async with httpx.AsyncClient(timeout=180) as client_http:
        r = await client_http.get(video_url)
        r.raise_for_status()
        return request_id, r.content


@dp.message(Command("img"))
async def cmd_img(message: Message):
    prompt = (message.text or "").replace("/img", "", 1).strip()
    if not prompt:
        await message.answer("Напиши так: /img описание картинки")
        return

    try:
        img_bytes = await run_with_thinking(message.bot, message.chat.id, generate_image_from_text_with_fallback(prompt))
        path = _save_bytes_to_tmp(f"img_{int(time.time())}.png", img_bytes)
        await message.answer_document(FSInputFile(path), caption="Готово ✅")
    except Exception as e:
        logging.exception("DeepAI img failed")
        await message.answer(f"Ошибка генерации: {e}")


@dp.message(Command("enhance"))
async def cmd_enhance(message: Message):
    user_id = message.from_user.id
    src = last_images.get(user_id)

    # scale: /enhance 2, /enhance 4, /enhance 2x, /enhance 4x
    m = re.search(r"\b(2|4)\s*x?\b", (message.text or ""))
    scale = int(m.group(1)) if m else ENHANCE_DEFAULT_SCALE

    if not src:
        await message.answer("Сначала пришли фото, которое нужно улучшить.")
        return

    try:
        img_bytes = await run_with_thinking(message.bot, message.chat.id, enhance_image_with_fallback(src, scale=scale))
        path = _save_bytes_to_tmp(f"enhanced_{int(time.time())}.png", img_bytes)
        await message.answer_document(FSInputFile(path), caption="Улучшил ✅")
    except Exception as e:
        logging.exception("DeepAI enhance failed")
        await message.answer(f"Ошибка улучшения: {e}")


@dp.message(Command("img2vid"))
async def cmd_img2vid(message: Message):
    """
    /img2vid <описание> — сделать видео из последнего присланного фото.
    """
    await _react_ok(message)
    user_id = message.from_user.id
    prompt = (message.text or "").replace("/img2vid", "", 1).strip()
    if not prompt:
        await message.answer("Напиши так: /img2vid описание (движение/камера/стиль). Или нажми 🎥 под фото.")
        return

    src = last_images.get(user_id)
    if not src:
        await message.answer("Сначала пришли фото, из которого нужно сделать видео.")
        return

    try:
        await message.bot.send_chat_action(message.chat.id, ChatAction.UPLOAD_VIDEO)
        task_id, video_bytes = await run_with_thinking(
            message.bot,
            message.chat.id,
            pika_image_bytes_to_video(src, prompt=prompt),
        )
        path = _save_bytes_to_tmp(f"pika_img2vid_{int(time.time())}.mp4", video_bytes)
        await message.answer_video(FSInputFile(path), caption="Видео готово ✅")
    except Exception as e:
        logging.exception("Runway img2vid failed")
        await message.answer(f"Ошибка генерации видео: {e}")


@dp.message(Command("txt2vid"))
async def cmd_txt2vid(message: Message):
    await _react_ok(message)
    await message.answer("txt2vid сейчас не подключен. Подключен только img2vid через Pika (fal.ai).")




@dp.message(Command("wb_retouch"))
async def cmd_wb_retouch(message: Message):
    user_id = message.from_user.id
    src = last_images.get(user_id)

    if not src:
        await message.answer("Сначала пришли фото, которое нужно отретушировать под WB.")
        return

    try:
        img_bytes = await run_with_thinking(callback.bot, callback.message.chat.id, retouch_for_wb(src))
        path = _save_bytes_to_tmp(f"wb_retouch_{int(time.time())}.png", img_bytes)
        await message.answer_document(FSInputFile(path), caption="WB ретушь готова ✅")
    except Exception as e:
        logging.exception("WB retouch failed")
        await message.answer(f"Ошибка ретуши: {e}")


@dp.callback_query(F.data == ENHANCE_CB)
async def cb_enhance_last(callback: CallbackQuery):
    user_id = callback.from_user.id
    src = last_images.get(user_id)
    if not src:
        await callback.message.answer("Не нашёл последнее фото. Пришли фото ещё раз.")
        await callback.answer()
        return
    await callback.answer("Улучшаю…")
    try:
        img_bytes = await enhance_image_with_fallback(src, scale=ENHANCE_DEFAULT_SCALE)
        path = _save_bytes_to_tmp(f"enhanced_{int(time.time())}.png", img_bytes)
        await callback.message.answer_document(FSInputFile(path), caption="Улучшил ✅")
    except Exception as e:
        logging.exception("Callback enhance failed")
        await callback.message.answer(f"Ошибка улучшения: {e}")


@dp.callback_query(F.data == RETROUCH_CB)
async def cb_wb_retouch_last(callback: CallbackQuery):
    user_id = callback.from_user.id
    src = last_images.get(user_id)
    if not src:
        await callback.message.answer("Не нашёл последнее фото. Пришли фото ещё раз.")
        await callback.answer()
        return
    await callback.answer("Ретуширую…")
    try:
        img_bytes = await run_with_thinking(callback.bot, callback.message.chat.id, retouch_for_wb(src))
        path = _save_bytes_to_tmp(f"wb_retouch_{int(time.time())}.png", img_bytes)
        await callback.message.answer_document(FSInputFile(path), caption="WB ретушь готова ✅")
    except Exception as e:
        logging.exception("Callback WB retouch failed")
        await callback.message.answer(f"Ошибка ретуши: {e}")



@dp.callback_query(F.data == VIDEO_CB)
async def cb_video_last(callback: CallbackQuery):
    user_id = callback.from_user.id
    if user_id not in last_images:
        await callback.message.answer("Не нашёл последнее фото. Пришли фото ещё раз.")
        await callback.answer()
        return

    pending_video_prompt[user_id] = True
    await callback.answer("Ок 👍")
    await callback.message.answer(
        "🎥 Принял! Теперь пришли ОДНИМ сообщением текст-описание для видео (движение/камера/стиль).\n"
        "Например: «Плавный наезд камеры, лёгкое вращение товара, студийный свет».\n\n"
        "Либо используй команду: /img2vid <текст>"
    )


# ---------- reactions ----------
async def _react_ok(message: Message):
    try:
        await message.bot.set_message_reaction(
            chat_id=message.chat.id,
            message_id=message.message_id,
            reaction=[ReactionTypeEmoji(emoji="👌")]
        )
    except Exception:
        pass

# ---------- thinking indicator ----------
async def _thinking_indicator(bot: Bot, chat_id: int, stop_event: asyncio.Event):
    """Показывает 'бот печатает…' пока не установлен stop_event."""
    try:
        while not stop_event.is_set():
            await bot.send_chat_action(chat_id, ChatAction.TYPING)
            await asyncio.sleep(4)
    except Exception:
        pass

async def run_with_thinking(bot: Bot, chat_id: int, coro):
    """Оборачивает долгую операцию с индикатором печати."""
    stop_event = asyncio.Event()
    task = asyncio.create_task(_thinking_indicator(bot, chat_id, stop_event))
    try:
        result = await coro
        return result
    finally:
        stop_event.set()
        await asyncio.sleep(0.1)
        task.cancel()

# ---------- commands ----------
@dp.message(Command("start"))
async def welcome(message: Message):
    await message.answer(
        "Привет, я Альтератти, бот, созданный thebkr.\n"
    )


@dp.message(Command("help"))
async def helper(message: Message):
    await message.answer(
        "Команды:\n"
        "/start — старт\n"
        "/clear — очистить контекст\n\n"
        "Редактирование файлов:\n"
        "1) отправь файл\n"
        "2) напиши: /edit <инструкция>\n\n"
        "Создание файлов:\n"
        "/make_xlsx <что должно быть в таблице>\n"
        "/make_docx <что должно быть в документе>\n"
        "/make_txt <что должно быть в тексте>\n\n"
        "Интернет-поиск без команд:\n"
        "Напиши в обычном сообщении: 'найди …', 'дай ссылки …', 'что сейчас …', 'актуальные правила …' — я сам поищу.\n\n"
        "Команды поиска (если всё же надо): /search и /research"
    )


@dp.message(Command("clear"))
async def clear(message: Message):
    clear_past()
    await message.answer("Ок, очистил контекст.")


# --- твои /edit /make_* /search /research и обработчики файлов ниже ОСТАЮТСЯ ---
# (Я их не вырезал — они у тебя уже есть.)


@dp.message(Command("edit"))
async def edit_last_file(message: Message, bot: Bot):
    user_id = message.from_user.id
    instructions = (message.text or "").replace("/edit", "", 1).strip()
    if not instructions:
        await message.answer("Напиши инструкцию так: /edit что именно поменять в последнем файле")
        return

    lf = last_files.get(user_id)
    if not lf:
        await message.answer("Сначала пришли файл, который нужно изменить.")
        return

    if len(lf.data) > MAX_FILE_BYTES:
        await message.answer("Файл слишком большой. Пришли поменьше (до ~12MB).")
        return

    ext = lf.ext
    try:
        if ext in {".txt", ".csv", ".json", ".md", ".log"}:
            src = _extract_text_from_plain(lf.data)
            edited = await _edit_text_like(src, instructions)
            out_bytes = edited.encode("utf-8")
            out_name = f"edited_{re.sub(r'[^a-zA-Z0-9._-]+','_', lf.filename)}"
            path = _save_bytes_to_tmp(out_name, out_bytes)
            await message.answer_document(FSInputFile(path), caption="Готово. Вот изменённый файл.")
            return

        if ext == ".docx":
            out_bytes, _ = await _edit_docx_bytes(lf.data, instructions)
            out_name = f"edited_{os.path.splitext(lf.filename)[0]}.docx"
            path = _save_bytes_to_tmp(out_name, out_bytes)
            await message.answer_document(FSInputFile(path), caption="Готово. Вот изменённый DOCX.")
            return

        if ext in {".xlsx", ".xlsm"}:
            out_bytes, _ = await _edit_xlsx_bytes(lf.data, instructions)
            out_name = f"edited_{os.path.splitext(lf.filename)[0]}.xlsx"
            path = _save_bytes_to_tmp(out_name, out_bytes)
            await message.answer_document(FSInputFile(path), caption="Готово. Вот изменённый XLSX.")
            return

        if ext == ".pdf":
            src = _extract_text_from_pdf(lf.data, max_pages=10)
            edited = await _edit_text_like(src, instructions)
            out_bytes = _build_docx_from_text(edited)
            out_name = f"edited_{os.path.splitext(lf.filename)[0]}.docx"
            path = _save_bytes_to_tmp(out_name, out_bytes)
            await message.answer_document(
                FSInputFile(path),
                caption="PDF как исходник — сложен для правок. Сделал DOCX-версию с изменениями."
            )
            return

        await message.answer("Этот тип файла пока не умею править. Пришли txt/docx/xlsx/pdf.")
    except Exception as e:
        logging.exception("Edit failed")
        await message.answer(f"Не смог применить правки. Ошибка: {e}")


@dp.message(Command("make_xlsx"))
async def make_xlsx(message: Message):
    prompt = (message.text or "").replace("/make_xlsx", "", 1).strip()
    if not prompt:
        await message.answer("Пример: /make_xlsx Сделай таблицу: SKU, Цена, CTR, CR (10 строк демо)")
        return

    system = (
        "Ты генерируешь таблицы. Верни ТОЛЬКО TSV (табами, строки переносами), без markdown и пояснений. "
        "Первая строка — заголовки."
    )
    user = f"Сгенерируй таблицу по запросу:\n{prompt}"
    tsv = await _ask_openai_text(system, user)
    tsv = tsv.strip().strip("```").strip()

    out_bytes = _build_xlsx_from_tsv(tsv)
    out_name = "generated.xlsx"
    path = _save_bytes_to_tmp(out_name, out_bytes)
    await message.answer_document(FSInputFile(path), caption="Сгенерировал XLSX.")


@dp.message(Command("make_docx"))
async def make_docx(message: Message):
    prompt = (message.text or "").replace("/make_docx", "", 1).strip()
    if not prompt:
        await message.answer("Пример: /make_docx Составь регламент обработки отзывов на WB на 1 страницу")
        return

    system = "Сгенерируй документ по запросу. Верни ТОЛЬКО чистый текст документа, без markdown."
    text = await _ask_openai_text(system, prompt)

    out_bytes = _build_docx_from_text(text)
    out_name = "generated.docx"
    path = _save_bytes_to_tmp(out_name, out_bytes)
    await message.answer_document(FSInputFile(path), caption="Сгенерировал DOCX.")


@dp.message(Command("make_txt"))
async def make_txt(message: Message):
    prompt = (message.text or "").replace("/make_txt", "", 1).strip()
    if not prompt:
        await message.answer("Пример: /make_txt Напиши 10 вариантов оффера для карточки товара")
        return

    system = "Сгенерируй текст по запросу. Верни ТОЛЬКО результат, без пояснений."
    text = await _ask_openai_text(system, prompt)

    out_bytes = (text or "").encode("utf-8")
    out_name = "generated.txt"
    path = _save_bytes_to_tmp(out_name, out_bytes)
    await message.answer_document(FSInputFile(path), caption="Сгенерировал TXT.")


@dp.message(Command("search"))
async def cmd_search(message: Message):
    q = (message.text or "").replace("/search", "", 1).strip()
    if not q:
        await message.answer("Напиши так: /search запрос")
        return

    try:
        results = await serper_search_cached(q, num=SEARCH_RESULTS_NUM)
    except Exception as e:
        await message.answer(f"Ошибка поиска: {e}")
        return

    if not results:
        await message.answer("Ничего не нашёл. Попробуй переформулировать запрос.")
        return

    await message.answer(format_search_results(results))


@dp.message(Command("research"))
async def cmd_research(message: Message):
    q = (message.text or "").replace("/research", "", 1).strip()
    if not q:
        await message.answer("Напиши так: /research запрос (я найду и сделаю краткий вывод)")
        return

    try:
        results = await serper_search_cached(q, num=SEARCH_RESULTS_NUM)
    except Exception as e:
        await message.answer(f"Ошибка поиска: {e}")
        return

    if not results:
        await message.answer("Ничего не нашёл. Попробуй переформулировать запрос.")
        return

    listing = format_search_results(results)
    await message.answer(listing)

    system = (
        "Ты делаешь краткую сводку по результатам интернет-поиска. "
        "Сначала 3-7 буллетов с выводами, затем 'Что сделать дальше' (3-5 шагов). "
        "Если фактов мало — скажи, чего не хватает."
    )
    user = "Запрос: " + q + "\n\n" + format_results_for_prompt(results)

    try:
        summary = await _ask_openai_text(system, user)
        await message.answer("🧠 Сводка:\n\n" + summary)
    except Exception as e:
        logging.exception("OpenAI summary failed")
        await message.answer(f"Не смог сделать сводку: {e}")


# ---------- file handlers ----------
@dp.message(F.photo)
async def handle_photo(message: Message, bot: Bot):
    await _react_ok(message)
    photo = message.photo[-1]
    image_bytes, _ = await _download_telegram_file(bot, photo.file_id)

    last_images[message.from_user.id] = image_bytes

    prompt = (
        "Распознай, что на изображении, и помоги пользователю.\n"
        "Если это документ/скрин — извлеки ключевой текст, найди ошибки/суть и дай рекомендации.\n"
    )
    try:
        answer = await run_with_thinking(message.bot, message.chat.id, _ask_openai_vision(prompt, image_bytes))
    except Exception as e:
        logging.exception("OpenAI vision failed")
        await message.answer(f"OpenAI error: {e}")
        return

    reference.response = answer
    await message.answer(answer)
    await message.answer("Что сделать с фото?", reply_markup=_image_action_keyboard())


@dp.message(F.document)
async def handle_document(message: Message, bot: Bot):
    await _react_ok(message)
    doc = message.document
    filename = doc.file_name or "file"
    ext = _ext(filename)
    mime = doc.mime_type or ""

    file_bytes, _ = await _download_telegram_file(bot, doc.file_id)

    if len(file_bytes) > MAX_FILE_BYTES:
        await message.answer("Файл слишком большой. Пришли поменьше (до ~12MB).")
        return

    last_files[message.from_user.id] = LastFile(
        filename=filename,
        ext=ext,
        mime=mime,
        data=file_bytes,
    )

    caption = (message.caption or "").strip()
    if caption.lower().startswith("edit:"):
        instructions = caption[5:].strip()
        message.text = f"/edit {instructions}"
        await edit_last_file(message, bot)
        return

    if ext in {".png", ".jpg", ".jpeg", ".webp"} or mime.startswith("image/"):
        last_images[message.from_user.id] = file_bytes
        prompt = f"Пользователь прислал изображение файлом: {filename}. Распознай и помоги."
        try:
            answer = await run_with_thinking(message.bot, message.chat.id, _ask_openai_vision(prompt, file_bytes))
        except Exception as e:
            logging.exception("OpenAI vision failed")
            await message.answer(f"OpenAI error: {e}")
            return
        reference.response = answer
        await message.answer(answer)
        await message.answer("Что сделать с изображением?", reply_markup=_image_action_keyboard())
        await message.answer("Файл запомнил. Если нужно изменить — напиши /edit <что поменять>.")
        return

    try:
        if ext == ".pdf":
            extracted = _extract_text_from_pdf(file_bytes, max_pages=10)
        elif ext == ".docx":
            extracted = _extract_text_from_docx(file_bytes)
        elif ext in {".xlsx", ".xlsm"}:
            extracted = _extract_tsv_preview_from_xlsx(file_bytes, max_rows=120, max_cols=25)
        else:
            extracted = _extract_text_from_plain(file_bytes)
    except Exception as e:
        logging.exception("File parse failed")
        await message.answer(f"Не смог прочитать файл {filename}. Ошибка: {e}")
        return

    extracted = _safe_truncate(extracted, max_chars=45_000)

    user_hint = caption if caption else "Разбери файл: что в нём важного, ошибки/риски, и что делать дальше."
    system = (
        "Отвечай по делу. Если это таблица — дай выводы и рекомендации. "
        "Если это документ — резюме и конкретные улучшения."
    )
    user = (
        f"Файл: {filename}\n"
        f"Задача: {user_hint}\n\n"
        f"Содержимое (может быть обрезано):\n-----\n{extracted}\n-----\n"
    )

    try:
        answer = await run_with_thinking(message.bot, message.chat.id, _ask_openai_text(system, user))
    except Exception as e:
        logging.exception("OpenAI text failed")
        await message.answer(f"OpenAI error: {e}")
        return

    reference.response = answer
    await message.answer(answer)
    await message.answer("Файл запомнил. Если нужно изменить — напиши /edit <что поменять>.")


@dp.message(F.text)
async def chat_gpt(message: Message):
    await _react_ok(message)
    user_text = (message.text or "").strip()
    if not user_text:
        return
    user_id = message.from_user.id

    # Если мы ждали промпт для видео (после кнопки 🎥)
    if pending_video_prompt.pop(user_id, False) and not user_text.startswith("/"):
        src = last_images.get(user_id)
        if not src:
            await message.answer("Не вижу последнее фото. Пришли фото ещё раз.")
            return
        try:
            await message.bot.send_chat_action(message.chat.id, ChatAction.UPLOAD_VIDEO)
            task_id, video_bytes = await run_with_thinking(
                message.bot,
                message.chat.id,
                pika_image_bytes_to_video(src, prompt=user_text),
            )
            path = _save_bytes_to_tmp(f"pika_video_{int(time.time())}.mp4", video_bytes)
            await message.answer_video(FSInputFile(path), caption="Видео готово ✅")
        except Exception as e:
            logging.exception("Runway img2vid failed (pending mode)")
            await message.answer(f"Ошибка генерации видео: {e}")
        return


    # Команды не трогаем
    if user_text.startswith("/"):
        return

    # --- AUT0SEARCH ---
    search_block = ""
    used_search = False

    if SERPER_API_KEY and should_autosearch(user_text) and can_search_now(message.from_user.id):
        try:
            results = await serper_search_cached(user_text, num=SEARCH_RESULTS_NUM)
            if results:
                used_search = True
                search_block = (
                    "\n\nРЕЗУЛЬТАТЫ ПОИСКА (snippets):\n"
                    + format_results_for_prompt(results)
                    + "\n\nТребование: используй результаты поиска, добавь 2–5 ссылок."
                )
        except Exception:
            logging.exception("Autosearch failed")
            # падаем обратно на обычный ответ без поиска

    system = (
        "Отвечай по делу, человеческим языком. "
        "Если есть блок 'РЕЗУЛЬТАТЫ ПОИСКА' — опирайся на него и добавь ссылки. "
        "Если данных недостаточно — скажи, чего не хватает и что уточнить."
    )

    # подмешиваем контекст + результаты поиска
    user = user_text + search_block

    try:
        answer = await run_with_thinking(message.bot, message.chat.id, _ask_openai_text(system, user))
    except Exception as e:
        logging.exception("OpenAI request failed")
        await message.answer(f"OpenAI error: {e}")
        return

    reference.response = answer

    # (опционально) маленький индикатор, что поиск применён
    if used_search:
        await message.answer("🌐 Нашёл в интернете, вот по сути:\n\n" + answer)
    else:
        await message.answer(answer)


# ---------- entrypoint ----------
async def main():
    bot = Bot(token=TELEGRAM_BOT_TOKEN)
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())